"""Render a markdown doc to a styled, reviewable .docx (Word) file.

Companion to md2pdf.py — same EY visual language (yellow rule under section
headings, Georgia body, Segoe UI headings, Consolas code) but built as real Word
structures rather than printed pixels: true Heading 1-4 styles so the Navigation
pane and any inserted TOC work, real tables you can sort and comment on, and the
hand-written contents list wired to real bookmarks so it is clickable in Word.

Usage:  python scripts/md2docx.py docs/REQUIREMENTS.md docs/REQUIREMENTS.docx
        python scripts/md2docx.py docs/REQUIREMENTS.md out.docx --page-break-sections
Needs:  pip install python-docx

--page-break-sections starts every "## " section on a fresh page (what md2pdf
does for print). Off by default: in Word the Navigation pane is the index, and
hard breaks fight reviewers who add comments and content.
"""
from __future__ import annotations

import re
import sys
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

# --- palette (matches md2pdf.py) -------------------------------------------
EY_YELLOW = "FFE600"
INK = RGBColor(0x16, 0x16, 0x1C)
INK_DARK = RGBColor(0x0D, 0x0D, 0x13)
MUTED = RGBColor(0x55, 0x55, 0x5F)
CODE_BG = "F3F3F7"
BLOCK_BG = "F7F7FA"
CALLOUT_BG = "FFFBE0"
CALLOUT_EDGE = "D9B800"
RULE = "DCDCE4"
TH_BG = "F3F3F7"
TH_EDGE = "C9C9D4"

BODY_FONT = "Georgia"
HEAD_FONT = "Segoe UI"
MONO_FONT = "Consolas"

# pPr children must appear in schema order; these are the elements that may
# legally follow a border/shading block.
_AFTER_PBDR = ("w:shd", "w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")
_AFTER_PSHD = ("w:tabs", "w:spacing", "w:ind", "w:jc", "w:rPr", "w:sectPr")
_AFTER_TCSHD = ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign", "w:hideMark")


# --------------------------------------------------------------------------
# low-level OOXML helpers (python-docx has no API for borders / shading)
# --------------------------------------------------------------------------
def _el(tag: str, **attrs) -> OxmlElement:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), str(v))
    return e


def _shade_para(par, fill: str) -> None:
    pPr = par._p.get_or_add_pPr()
    pPr.insert_element_before(_el("w:shd", val="clear", color="auto", fill=fill), *_AFTER_PSHD)


def _shade_cell(cell, fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.insert_element_before(_el("w:shd", val="clear", color="auto", fill=fill), *_AFTER_TCSHD)


def _border_para(par, **sides) -> None:
    """sides: left=(sz_eighths_pt, color) etc. sz is in 1/8 pt."""
    pPr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        if side in sides:
            sz, color = sides[side]
            bdr.append(_el(f"w:{side}", val="single", sz=sz, space=4, color=color))
    pPr.insert_element_before(bdr, *_AFTER_PBDR)


def _shade_run(run, fill: str) -> None:
    run._r.get_or_add_rPr().append(_el("w:shd", val="clear", color="auto", fill=fill))


def _table_borders(table, color: str = RULE) -> None:
    """Horizontal rules only — matches the PDF's table styling."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "bottom", "insideH"):
        borders.append(_el(f"w:{side}", val="single", sz=4, space=0, color=color))
    for side in ("left", "right", "insideV"):
        borders.append(_el(f"w:{side}", val="none", sz=0, space=0, color="auto"))
    tblPr.append(borders)


def _row_keep_together(row, header: bool = False) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(_el("w:cantSplit", val="true"))
    if header:
        trPr.append(_el("w:tblHeader", val="true"))


def _bookmark(par, name: str, bid: int) -> None:
    p = par._p
    idx = 1 if p.find(qn("w:pPr")) is not None else 0
    p.insert(idx, _el("w:bookmarkStart", id=bid, name=name))
    p.append(_el("w:bookmarkEnd", id=bid))


# --------------------------------------------------------------------------
# inline markdown -> runs
# --------------------------------------------------------------------------
_TOKEN = re.compile(
    r"`(?P<code>[^`]+)`"
    r"|\[(?P<ltext>[^\]]+)\]\((?P<lhref>[^)]+)\)"
    r"|\*\*(?P<bold>.+?)\*\*"
    r"|(?<![\*\w])\*(?P<ital>[^*\n]+)\*(?!\*)"
)


@dataclass
class Tok:
    text: str
    bold: bool = False
    ital: bool = False
    code: bool = False
    href: str | None = None


def tokenize(text: str, bold=False, ital=False, href=None) -> list[Tok]:
    out: list[Tok] = []
    pos = 0
    for m in _TOKEN.finditer(text):
        if m.start() > pos:
            out.append(Tok(text[pos:m.start()], bold, ital, False, href))
        if m.group("code") is not None:
            out.append(Tok(m.group("code"), bold, ital, True, href))
        elif m.group("ltext") is not None:
            out += tokenize(m.group("ltext"), bold, ital, m.group("lhref"))
        elif m.group("bold") is not None:
            out += tokenize(m.group("bold"), True, ital, href)
        else:
            out += tokenize(m.group("ital"), bold, True, href)
        pos = m.end()
    if pos < len(text):
        out.append(Tok(text[pos:], bold, ital, False, href))
    return out


def unescape(s: str) -> str:
    return re.sub(r"\\([\\`*_{}\[\]()#+\-.!|])", r"\1", s)


class Renderer:
    def __init__(self, doc: Document, page_break_sections: bool):
        self.doc = doc
        self.page_break_sections = page_break_sections
        self.bookmarks: dict[str, str] = {}   # slug -> word bookmark name
        self._bid = 0
        self._seen_h2 = False

    # -- anchors ----------------------------------------------------------
    @staticmethod
    def slugify(text: str) -> str:
        """GitHub's anchor style, so the hand-written contents links resolve.

        Each whitespace char becomes its own hyphen — GitHub does NOT collapse
        runs, so "features — roadmap" (em dash stripped) anchors as
        "features--roadmap". Collapsing here silently breaks those links.
        """
        s = re.sub(r"[^\w\s-]", "", text.strip().lower())
        return re.sub(r"\s", "-", s).strip("-")

    def bookmark_name(self, slug: str) -> str:
        """Word bookmark names: letters/digits/underscore, must not start with a
        digit, 40 chars max."""
        if slug in self.bookmarks:
            return self.bookmarks[slug]
        base = "s_" + re.sub(r"[^\w]", "_", slug)[:34]
        name, n = base, 2
        while name in self.bookmarks.values():
            name, n = f"{base}_{n}", n + 1
        self.bookmarks[slug] = name
        return name

    def next_bid(self) -> int:
        self._bid += 1
        return self._bid

    # -- runs -------------------------------------------------------------
    def emit(self, par, tokens: list[Tok], size: Pt | None = None) -> None:
        for t in tokens:
            text = unescape(t.text)
            if not text:
                continue
            run = par.add_run(text)
            run.bold = t.bold or None
            run.italic = t.ital or None
            if t.code:
                run.font.name = MONO_FONT
                run.font.size = Pt(8.5) if size is None else Pt(size.pt - 1.6)
                _shade_run(run, CODE_BG)
            elif size is not None:
                run.font.size = size
            if t.href:
                run.font.color.rgb = RGBColor(0x1A, 0x4C, 0x8C)
                run.underline = True
                self._wrap_hyperlink(par, run, t.href)

    def _wrap_hyperlink(self, par, run, href: str) -> None:
        link = OxmlElement("w:hyperlink")
        if href.startswith("#"):
            link.set(qn("w:anchor"), self.bookmark_name(href[1:]))
        else:
            rid = par.part.relate_to(href, RT.HYPERLINK, is_external=True)
            link.set(qn("r:id"), rid)
        par._p.remove(run._r)
        link.append(run._r)
        par._p.append(link)

    # -- blocks -----------------------------------------------------------
    def heading(self, level: int, text: str) -> None:
        par = self.doc.add_paragraph(style=f"Heading {level}")
        if level == 2:
            if self.page_break_sections and self._seen_h2:
                par.paragraph_format.page_break_before = True
            self._seen_h2 = True
            _border_para(par, bottom=(16, EY_YELLOW))
        self.emit(par, tokenize(text))
        _bookmark(par, self.bookmark_name(self.slugify(re.sub(r"[*`]", "", text))), self.next_bid())

    def paragraph(self, text: str) -> None:
        par = self.doc.add_paragraph()
        self.emit(par, tokenize(text))

    def bullet(self, text: str, ordered: bool, indent: int) -> None:
        par = self.doc.add_paragraph(style="List Number" if ordered else "List Bullet")
        par.paragraph_format.left_indent = Mm(6 + 6 * indent)
        self.emit(par, tokenize(text))

    def code_block(self, lines: list[str]) -> None:
        par = self.doc.add_paragraph()
        pf = par.paragraph_format
        pf.space_before, pf.space_after = Pt(6), Pt(8)
        pf.left_indent = Mm(2)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.keep_together = True
        _shade_para(par, BLOCK_BG)
        _border_para(par, left=(20, EY_YELLOW), top=(4, RULE), bottom=(4, RULE), right=(4, RULE))
        for i, line in enumerate(lines):
            if i:
                par.add_run().add_break()
            run = par.add_run(line)
            run.font.name = MONO_FONT
            run.font.size = Pt(7.5)

    def callout(self, lines: list[str]) -> None:
        text = " ".join(l.strip() for l in lines if l.strip())
        par = self.doc.add_paragraph()
        pf = par.paragraph_format
        pf.space_before, pf.space_after = Pt(6), Pt(8)
        pf.left_indent, pf.right_indent = Mm(3), Mm(2)
        pf.keep_together = True
        _shade_para(par, CALLOUT_BG)
        _border_para(par, left=(18, CALLOUT_EDGE), top=(4, CALLOUT_EDGE),
                     bottom=(4, CALLOUT_EDGE), right=(4, CALLOUT_EDGE))
        self.emit(par, tokenize(text), size=Pt(9.5))

    def rule(self) -> None:
        par = self.doc.add_paragraph()
        par.paragraph_format.space_before = Pt(4)
        par.paragraph_format.space_after = Pt(8)
        _border_para(par, bottom=(6, RULE))

    def table(self, rows: list[list[str]]) -> None:
        head, body = rows[0], rows[1:]
        table = self.doc.add_table(rows=len(rows), cols=len(head))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True
        _table_borders(table)

        for j, cell_text in enumerate(head):
            cell = table.cell(0, j)
            _shade_cell(cell, TH_BG)
            par = cell.paragraphs[0]
            par.paragraph_format.space_before = Pt(2)
            par.paragraph_format.space_after = Pt(2)
            run = par.add_run(re.sub(r"[*`]", "", cell_text).upper())
            run.font.name, run.font.size, run.bold = HEAD_FONT, Pt(7), True
            run.font.color.rgb = MUTED
        _row_keep_together(table.rows[0], header=True)

        for i, row in enumerate(body, start=1):
            for j in range(len(head)):
                cell = table.cell(i, j)
                par = cell.paragraphs[0]
                par.paragraph_format.space_before = Pt(2)
                par.paragraph_format.space_after = Pt(2)
                par.paragraph_format.line_spacing = 1.1
                text = row[j] if j < len(row) else ""
                self.emit(par, tokenize(text), size=Pt(8.2))
                for r in par.runs:
                    if r.font.name != MONO_FONT:
                        r.font.name = HEAD_FONT
            _row_keep_together(table.rows[i])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)


# --------------------------------------------------------------------------
# block-level markdown parser
# --------------------------------------------------------------------------
_SPLIT_CELLS = re.compile(r"(?<!\\)\|")


def split_row(line: str) -> list[str]:
    parts = _SPLIT_CELLS.split(line.strip())
    if parts and not parts[0].strip():
        parts = parts[1:]
    if parts and not parts[-1].strip():
        parts = parts[:-1]
    return [p.strip() for p in parts]


def is_divider(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:|-]*-[\s:|-]*\|?", line.strip())) and "-" in line and "|" in line


def render(md: str, doc: Document, page_break_sections: bool) -> None:
    r = Renderer(doc, page_break_sections)
    lines = md.splitlines()
    i, n = 0, len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # fenced code
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            r.code_block(buf)
            continue

        # heading
        if m := re.match(r"(#{1,4})\s+(.*)", stripped):
            r.heading(len(m.group(1)), m.group(2).strip())
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            r.rule()
            i += 1
            continue

        # blockquote / callout
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(re.sub(r"^\s*>\s?", "", lines[i]))
                i += 1
            r.callout(buf)
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_divider(lines[i + 1]):
            head = split_row(lines[i])
            i += 2
            rows = [head]
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            r.table(rows)
            continue

        # list item
        if m := re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)", line):
            indent = len(m.group(1)) // 2
            ordered = m.group(2)[0].isdigit()
            body = [m.group(3)]
            i += 1
            # fold continuation lines into the same item
            while i < n and lines[i].strip() and not re.match(
                r"^(\s*)([-*+]|\d+\.)\s+|^\s*(#{1,4}\s|>|\||```|-{3,}$)", lines[i]
            ):
                body.append(lines[i].strip())
                i += 1
            r.bullet(" ".join(body), ordered, indent)
            continue

        # paragraph
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^\s*(#{1,4}\s|>|\||```|[-*+]\s|\d+\.\s|-{3,}$)", lines[i]
        ):
            buf.append(lines[i].strip())
            i += 1
        r.paragraph(" ".join(buf))


# --------------------------------------------------------------------------
# document chrome
# --------------------------------------------------------------------------
def build_document(title: str) -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)          # A4
    sec.left_margin = sec.right_margin = Mm(18)
    sec.top_margin, sec.bottom_margin = Mm(18), Mm(18)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sizes = {1: 21, 2: 14, 3: 11.5, 4: 10.5}
    for level, size in sizes.items():
        st = doc.styles[f"Heading {level}"]
        st.font.name = HEAD_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = INK_DARK
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.space_before = Pt(0 if level == 1 else 14 - 2 * level)
        st.paragraph_format.space_after = Pt(4 if level > 2 else 6)

    for name in ("List Bullet", "List Number"):
        st = doc.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(10.5)
        st.font.color.rgb = INK
        st.paragraph_format.space_after = Pt(3)
        st.paragraph_format.line_spacing = 1.2

    # footer: page number
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    run.font.name, run.font.size, run.font.color.rgb = HEAD_FONT, Pt(8), MUTED
    fld_begin = _el("w:fldChar", fldCharType="begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = _el("w:fldChar", fldCharType="end")
    for e in (fld_begin, instr, fld_end):
        run._r.append(e)

    doc.core_properties.title = title
    return doc


def main() -> int:
    if len(sys.argv) < 3:
        print(__doc__)
        return 2
    src, out = Path(sys.argv[1]), Path(sys.argv[2])
    page_breaks = "--page-break-sections" in sys.argv[3:]

    md = src.read_text(encoding="utf-8")
    first_h1 = next((l[2:].strip() for l in md.splitlines() if l.startswith("# ")), src.stem)

    doc = build_document(first_h1)
    render(md, doc, page_breaks)
    doc.save(out)
    print(f"wrote {out}  ({out.stat().st_size / 1024:.0f} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
